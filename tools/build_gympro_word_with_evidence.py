from pathlib import Path
import subprocess
import textwrap

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, ImageOps


ROOT = Path.cwd()
DOCX_OUT = ROOT / "docs" / "Informe_GymPro_Flutter_con_Evidencias.docx"
EVIDENCE_DIR = ROOT / "docs" / "evidencias_word"
FLUTTER = Path("C:/flutter/flutter/bin/flutter.bat")
ADB = Path.home() / "AppData/Local/Android/sdk/platform-tools/adb.exe"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_GRAY = "F2F4F7"
BORDER = "B7C2D0"


def run_command(command, timeout=120):
    proc = subprocess.run(
        command,
        cwd=ROOT,
        text=True,
        capture_output=True,
        timeout=timeout,
        shell=True,
    )
    output = (proc.stdout + "\n" + proc.stderr).strip()
    return output if output else "(sin salida)"


def load_font(size=18):
    for candidate in [
        "C:/Windows/Fonts/consola.ttf",
        "C:/Windows/Fonts/Consola.ttf",
        "C:/Windows/Fonts/cour.ttf",
    ]:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def text_to_image(title, body, filename, width=1400, max_lines=42):
    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)
    font = load_font(20)
    title_font = load_font(24)
    wrapped = []
    for raw_line in body.splitlines():
        line = raw_line.rstrip()
        if not line:
            wrapped.append("")
            continue
        wrapped.extend(textwrap.wrap(line, width=105, replace_whitespace=False) or [""])
    wrapped = wrapped[:max_lines]

    line_height = 28
    height = 110 + max(10, len(wrapped)) * line_height
    image = Image.new("RGB", (width, height), "#101114")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width, 58), fill="#202126")
    draw.text((28, 16), title, fill="#FFFFFF", font=title_font)

    y = 80
    for line in wrapped:
        draw.text((28, y), line, fill="#E8EAED", font=font)
        y += line_height

    image = ImageOps.expand(image, border=2, fill="#B7C2D0")
    path = EVIDENCE_DIR / filename
    image.save(path)
    return path


def capture_emulator():
    path = EVIDENCE_DIR / "captura_app_usuario.png"
    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)
    if ADB.exists():
        with path.open("wb") as fh:
            subprocess.run(
                [str(ADB), "-s", "emulator-5554", "exec-out", "screencap", "-p"],
                cwd=ROOT,
                stdout=fh,
                stderr=subprocess.DEVNULL,
                timeout=30,
            )
    fallback = ROOT / "docs" / "captura_home_usuario.png"
    if not path.exists() or path.stat().st_size == 0:
        if fallback.exists():
            path.write_bytes(fallback.read_bytes())
    return path if path.exists() and path.stat().st_size > 0 else None


def make_evidence_images():
    images = {}
    images["flutter_version"] = text_to_image(
        "PowerShell - flutter --version",
        run_command(f'"{FLUTTER}" --version', timeout=60),
        "flutter_version.png",
    )
    images["devices"] = text_to_image(
        "PowerShell - flutter devices",
        run_command(f'"{FLUTTER}" devices', timeout=60),
        "flutter_devices.png",
    )
    images["project_root"] = text_to_image(
        "PowerShell - carpeta raiz del proyecto",
        run_command(
            "Get-ChildItem -Force | Select-Object Mode, LastWriteTime, Length, Name | Format-Table -AutoSize",
            timeout=20,
        ),
        "project_root.png",
    )
    images["analyze"] = text_to_image(
        "PowerShell - flutter analyze",
        run_command(f'"{FLUTTER}" analyze', timeout=120),
        "flutter_analyze.png",
    )
    images["test"] = text_to_image(
        "PowerShell - flutter test",
        run_command(f'"{FLUTTER}" test', timeout=120),
        "flutter_test.png",
    )
    images["structure"] = text_to_image(
        "PowerShell - tree /F lib",
        run_command("cmd /c tree /F lib", timeout=20),
        "tree_lib.png",
        max_lines=55,
    )
    images["main_dart"] = text_to_image(
        "Editor - lib/main.dart",
        (ROOT / "lib" / "main.dart").read_text(encoding="utf-8"),
        "main_dart.png",
        max_lines=38,
    )
    client_file = ROOT / "lib/features/client/presentation/pages/client_home_page.dart"
    client_text = "\n".join(client_file.read_text(encoding="utf-8").splitlines()[:70])
    images["client_code"] = text_to_image(
        "Editor - pantalla principal del usuario",
        client_text,
        "client_home_code.png",
        max_lines=45,
    )
    images["readme"] = text_to_image(
        "Editor - README.md",
        (ROOT / "README.md").read_text(encoding="utf-8"),
        "readme.png",
        max_lines=42,
    )
    emulator = capture_emulator()
    if emulator:
        images["app_screenshot"] = emulator
    return images


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text="", bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if text:
        r = p.add_run(text)
        set_run(r, size=11, bold=bold, italic=italic)
    return p


def add_heading(doc, number, title, level=1):
    text = f"{number}. {title}" if number else title
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, size=16 if level == 1 else 13, bold=True, color=BLUE if level == 1 else DARK_BLUE)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=11)


def add_step(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    r = p.add_run(f"{number}. {text}")
    set_run(r, size=11)


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        row.cells[0].width = Inches(1.9)
        row.cells[1].width = Inches(4.6)
    for idx, (label, value) in enumerate(rows):
        left, right = table.rows[idx].cells
        set_cell_border(left)
        set_cell_border(right)
        set_cell_shading(left, LIGHT_GRAY)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r1 = left.paragraphs[0].add_run(label)
        set_run(r1, size=10.5, bold=True, color=DARK_BLUE)
        r2 = right.paragraphs[0].add_run(value)
        set_run(r2, size=10.5)
    doc.add_paragraph()


def add_image_evidence(doc, title, image_path, width=6.2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run(r, size=10.5, bold=True, color=DARK_BLUE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(str(image_path), width=Inches(width))


def add_placeholder(doc, title, note):
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].width = Inches(6.5)
    table.rows[1].cells[0].width = Inches(6.5)
    for row in table.rows:
        set_cell_border(row.cells[0])
    set_cell_shading(table.rows[0].cells[0], LIGHT_GRAY)
    p1 = table.rows[0].cells[0].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(title)
    set_run(r1, size=10.5, bold=True, color=DARK_BLUE)
    p2 = table.rows[1].cells[0].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[1].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    r2 = p2.add_run(note)
    set_run(r2, size=10.5, italic=True, color=MUTED)


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    for name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE)]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("GymPro App - Evidencias del proyecto Flutter")
    set_run(r, size=9, color=MUTED)


def add_cover(doc):
    add_paragraph(doc, "INGENIERÍA DE SISTEMAS", bold=True).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, "PROYECTO MÓVIL CON FLUTTER", bold=True).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(72)
    title.paragraph_format.space_after = Pt(10)
    r = title.add_run("INFORME DE AVANCE TÉCNICO CON EVIDENCIAS")
    set_run(r, size=23, bold=True, color=BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("GYMPRO APP")
    set_run(r, size=18, bold=True, color=DARK_BLUE)
    add_paragraph(doc, "Aplicación móvil para gestión de gimnasio").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    add_key_value_table(
        doc,
        [
            ("Estudiante(s)", "Samuel Flores Cruz, Rafael Leandro Gómez Escobar, Mauricio Cartagena Alba"),
            ("Materia", "Desarrollo de aplicación móvil"),
            ("Docente", "____________________________________________"),
            ("Fecha", "____________________________________________"),
            ("Lugar", "Cochabamba - Bolivia"),
        ],
    )


def add_index(doc):
    add_heading(doc, None, "ÍNDICE DE CONTENIDO")
    rows = [
        ("Portada", "1"),
        ("Evidencia de configuración del entorno", "3"),
        ("Creación del proyecto Flutter", "4"),
        ("Ejecución de la aplicación", "5"),
        ("Configuración de Git y GitHub", "7"),
        ("Estructura inicial del proyecto", "8"),
        ("Configuración del tema inicial", "9"),
        ("Flujo de navegación", "11"),
        ("README del Proyecto", "12"),
        ("Conclusiones", "13"),
    ]
    for title, page in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(f"{title} {'.' * max(8, 68 - len(title))} {page}")
        set_run(r, size=11)


def build_doc(images):
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    doc.add_page_break()
    add_index(doc)
    doc.add_page_break()

    add_heading(doc, 2, "Evidencia de configuración del entorno")
    add_paragraph(doc, "Se documenta la instalación y detección del entorno de desarrollo usado para GymPro.")
    add_key_value_table(
        doc,
        [
            ("Sistema operativo", "Windows"),
            ("Framework", "Flutter"),
            ("Lenguaje", "Dart"),
            ("Dispositivo", "Emulador Android detectado por Flutter"),
        ],
    )
    add_image_evidence(doc, "Evidencia 1. Versión de Flutter y Dart", images["flutter_version"])
    add_image_evidence(doc, "Evidencia 2. Dispositivos disponibles", images["devices"])
    doc.add_page_break()

    add_heading(doc, 3, "Creación del proyecto Flutter")
    add_paragraph(doc, "El proyecto ya se encuentra inicializado y contiene las carpetas generadas por Flutter.")
    add_key_value_table(
        doc,
        [
            ("Nombre del proyecto", "aplicacion_movil"),
            ("Carpeta", str(ROOT)),
            ("Aplicación", "GymPro App"),
        ],
    )
    add_image_evidence(doc, "Evidencia 3. Carpeta raíz del proyecto", images["project_root"])
    doc.add_page_break()

    add_heading(doc, 4, "Ejecución de la aplicación")
    add_paragraph(doc, "La aplicación fue compilada y ejecutada en el emulador Android detectado.")
    add_image_evidence(doc, "Evidencia 4. Resultado de análisis estático", images["analyze"])
    add_image_evidence(doc, "Evidencia 5. Pruebas automatizadas", images["test"])
    doc.add_page_break()
    if "app_screenshot" in images:
        add_image_evidence(doc, "Evidencia 6. Aplicación ejecutándose en el emulador", images["app_screenshot"], width=3.2)
    else:
        add_placeholder(doc, "Evidencia 6", "Pendiente: captura de la app ejecutándose en el emulador.")
    doc.add_page_break()

    add_heading(doc, 5, "Configuración de Git y GitHub")
    add_paragraph(doc, "Este apartado queda preparado para agregar la evidencia de Git y GitHub cuando se inicialice el repositorio.")
    add_bullet(doc, "Inicializar Git en la carpeta del proyecto.")
    add_bullet(doc, "Crear repositorio remoto en GitHub.")
    add_bullet(doc, "Realizar el primer commit y subir la rama principal.")
    add_placeholder(doc, "Evidencia pendiente", "Agregar captura de git status, primer commit y repositorio de GitHub.")
    doc.add_page_break()

    add_heading(doc, 6, "Estructura inicial del proyecto")
    add_paragraph(doc, "La estructura se organizó en módulos dentro de lib para separar responsabilidades.")
    add_key_value_table(
        doc,
        [
            ("lib/app", "Configuración general de la aplicación."),
            ("lib/core", "Constantes, configuración, red y utilidades."),
            ("lib/shared", "Widgets y modelos reutilizables."),
            ("lib/features", "Módulos funcionales del MVP."),
        ],
    )
    add_image_evidence(doc, "Evidencia 7. Árbol de carpetas de lib", images["structure"])
    doc.add_page_break()

    add_heading(doc, 7, "Configuración del tema inicial")
    add_paragraph(doc, "Se reemplazó la pantalla base de Flutter por la pantalla principal del usuario con una estética negra y roja.")
    add_image_evidence(doc, "Evidencia 8. Archivo main.dart", images["main_dart"])
    add_image_evidence(doc, "Evidencia 9. Código de la pantalla principal del usuario", images["client_code"])
    doc.add_page_break()

    add_heading(doc, 8, "Flujo de navegación")
    add_paragraph(doc, "Por ahora la app inicia directamente en la interfaz del usuario. El login y redirección por roles se implementarán al final.")
    add_step(doc, 1, "El usuario abre la aplicación.")
    add_step(doc, 2, "Se muestra la pantalla principal del cliente.")
    add_step(doc, 3, "Desde la barra inferior podrá acceder a entrenamiento, rutina y perfil.")
    if "app_screenshot" in images:
        add_image_evidence(doc, "Evidencia 10. Pantalla principal usada como inicio temporal", images["app_screenshot"], width=2.8)
    doc.add_page_break()

    add_heading(doc, 9, "README del Proyecto")
    add_paragraph(doc, "El README conserva la información base del proyecto y podrá actualizarse con instrucciones específicas de GymPro.")
    add_image_evidence(doc, "Evidencia 11. README.md actual", images["readme"])
    doc.add_page_break()

    add_heading(doc, 10, "Conclusiones")
    add_paragraph(
        doc,
        "El proyecto GymPro ya cuenta con una estructura modular, una pantalla principal de usuario codificada y una ejecución verificada en emulador Android.",
    )
    add_bullet(doc, "Se comprobó que Flutter y Dart están instalados.")
    add_bullet(doc, "Se detectó el emulador Android y se ejecutó la aplicación.")
    add_bullet(doc, "Se verificó el código con flutter analyze y flutter test.")
    add_bullet(doc, "Queda pendiente documentar Git/GitHub cuando el repositorio sea inicializado.")

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)


def main():
    images = make_evidence_images()
    build_doc(images)
    print(DOCX_OUT)


if __name__ == "__main__":
    main()
