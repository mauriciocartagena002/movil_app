from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/Informe_GymPro_Flutter.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_GRAY = "F2F4F7"
BORDER = "B7C2D0"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
      borders = OxmlElement("w:tcBorders")
      tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def set_run(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text="", style=None, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if text:
        r = p.add_run(text)
        set_run(r, size=11, bold=bold, italic=italic)
    return p


def add_heading(doc, number, title, level=1):
    text = f"{number}. {title}" if number else title
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, size=16 if level == 1 else 13, bold=True, color=BLUE if level == 1 else DARK_BLUE)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=11)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=11)


def add_step(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    r = p.add_run(f"{number}. {text}")
    set_run(r, size=11)


def add_capture_box(doc, title, suggested_capture, height=2.0):
    table = doc.add_table(rows=2, cols=1)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        row.cells[0].width = Inches(6.5)
        set_cell_border(row.cells[0])
    set_cell_shading(table.rows[0].cells[0], LIGHT_GRAY)
    table.rows[0].height = Inches(0.35)
    table.rows[1].height = Inches(height)
    header = table.rows[0].cells[0].paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(title)
    set_run(run, size=10.5, bold=True, color=DARK_BLUE)
    body = table.rows[1].cells[0].paragraphs[0]
    body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[1].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    run = body.add_run(f"Espacio para captura: {suggested_capture}")
    set_run(run, size=10.5, italic=True, color=MUTED)


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, [1.9, 4.6])
    for idx, (label, value) in enumerate(rows):
        left, right = table.rows[idx].cells
        set_cell_border(left)
        set_cell_border(right)
        set_cell_shading(left, LIGHT_GRAY)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p1 = left.paragraphs[0]
        r1 = p1.add_run(label)
        set_run(r1, size=10.5, bold=True, color=DARK_BLUE)
        p2 = right.paragraphs[0]
        r2 = p2.add_run(value)
        set_run(r2, size=10.5)
    doc.add_paragraph()


def add_index(doc):
    add_heading(doc, None, "ÍNDICE DE CONTENIDO", 1)
    rows = [
        ("Portada", "1"),
        ("Evidencia de configuración del entorno", "3"),
        ("Creación del proyecto Flutter", "4"),
        ("Ejecución de la aplicación", "5"),
        ("Configuración de Git y GitHub", "6"),
        ("Estructura inicial del proyecto", "7"),
        ("Configuración del tema inicial", "8"),
        ("Flujo de navegación", "9"),
        ("README del Proyecto", "10"),
        ("Conclusiones", "11"),
    ]
    for title, page in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        line = f"{title} {'.' * max(8, 68 - len(title))} {page}"
        run = p.add_run(line)
        set_run(run, size=11)


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_after = Pt(6)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("GymPro App - Proyecto Flutter")
    set_run(run, size=9, color=MUTED)


def add_cover(doc):
    add_paragraph(doc, "INGENIERÍA DE SISTEMAS", bold=True).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, "PROYECTO MÓVIL CON FLUTTER", bold=True).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("INFORME DE AVANCE TÉCNICO")
    set_run(run, size=24, bold=True, color=BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("GYMPRO APP")
    set_run(run, size=18, bold=True, color=DARK_BLUE)
    add_paragraph(doc, "Aplicación móvil para gestión de gimnasio").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    add_key_value_table(
        doc,
        [
            ("Estudiante(s)", "Samuel Flores Cruz, Rafael Leandro Gómez Escobar, Mauricio Cartagena Alba"),
            ("Materia", "Desarrollo de aplicación móvil"),
            ("Docente", "____________________________________________"),
            ("Fecha", "____________________________________________"),
            ("Lugar", "Cochabamba - Bolivia"),
        ],
    )
    add_capture_box(doc, "Evidencia opcional", "logo del proyecto, portada de Figma o imagen representativa de GymPro", 1.5)


def add_environment(doc):
    add_heading(doc, 2, "Evidencia de configuración del entorno")
    add_paragraph(
        doc,
        "En esta sección se documenta la preparación del entorno de desarrollo utilizado para crear la aplicación móvil GymPro con Flutter y Dart.",
    )
    add_key_value_table(
        doc,
        [
            ("Sistema operativo", "Windows"),
            ("Framework", "Flutter"),
            ("Lenguaje", "Dart"),
            ("Editor sugerido", "Visual Studio Code o Android Studio"),
            ("Herramientas de apoyo", "Git, GitHub, emulador Android o dispositivo físico"),
        ],
    )
    add_bullet(doc, "Verificar que Flutter SDK esté instalado correctamente.")
    add_bullet(doc, "Verificar la instalación de Dart y Android SDK.")
    add_bullet(doc, "Comprobar el estado del entorno con el comando flutter doctor.")
    add_capture_box(doc, "Captura 1", "resultado del comando flutter doctor", 1.8)
    add_capture_box(doc, "Captura 2", "versiones de Flutter y Dart instaladas", 1.3)


def add_project_creation(doc):
    add_heading(doc, 3, "Creación del proyecto Flutter")
    add_paragraph(
        doc,
        "El proyecto fue inicializado como una aplicación Flutter, lo que permite desarrollar una solución móvil multiplataforma usando Dart como lenguaje principal.",
    )
    add_step(doc, 1, "Abrir una terminal en la carpeta donde se almacenará el proyecto.")
    add_step(doc, 2, "Ejecutar el comando de creación del proyecto Flutter.")
    add_step(doc, 3, "Abrir el proyecto generado en el editor de código.")
    add_key_value_table(
        doc,
        [
            ("Nombre del proyecto", "aplicacion_movil"),
            ("Nombre de la app", "GymPro"),
            ("Carpeta del proyecto", "C:\\Users\\user\\Documents\\proyectos\\aplicacion_movil"),
        ],
    )
    add_capture_box(doc, "Captura 3", "terminal con la creación del proyecto Flutter", 1.7)
    add_capture_box(doc, "Captura 4", "proyecto abierto en el editor", 1.7)


def add_execution(doc):
    add_heading(doc, 4, "Ejecución de la aplicación")
    add_paragraph(
        doc,
        "Después de crear el proyecto, se debe ejecutar la aplicación para comprobar que Flutter compila correctamente y que la pantalla inicial se muestra en un dispositivo o emulador.",
    )
    add_bullet(doc, "Seleccionar un emulador Android o conectar un teléfono físico.")
    add_bullet(doc, "Ejecutar la aplicación con flutter run.")
    add_bullet(doc, "Verificar que la pantalla inicial se abra sin errores.")
    add_capture_box(doc, "Captura 5", "aplicación ejecutándose en emulador o teléfono", 2.1)
    add_capture_box(doc, "Captura 6", "terminal mostrando la ejecución exitosa", 1.4)


def add_git(doc):
    add_heading(doc, 5, "Configuración de Git y GitHub")
    add_paragraph(
        doc,
        "La gestión del código fuente se realizará con Git y GitHub para mantener historial de cambios, respaldo del proyecto y colaboración entre integrantes.",
    )
    add_step(doc, 1, "Inicializar Git en la carpeta del proyecto.")
    add_step(doc, 2, "Crear el repositorio remoto en GitHub.")
    add_step(doc, 3, "Vincular el repositorio local con GitHub.")
    add_step(doc, 4, "Realizar el primer commit con la estructura inicial.")
    add_key_value_table(
        doc,
        [
            ("Repositorio local", "Pendiente de inicialización o verificación"),
            ("Repositorio GitHub", "____________________________________________"),
            ("Rama principal", "main"),
        ],
    )
    add_capture_box(doc, "Captura 7", "comandos git init, git status y primer commit", 1.7)
    add_capture_box(doc, "Captura 8", "repositorio creado en GitHub", 1.7)


def add_structure(doc):
    add_heading(doc, 6, "Estructura inicial del proyecto")
    add_paragraph(
        doc,
        "La estructura inicial se organizó por módulos para separar las responsabilidades de la aplicación y facilitar el crecimiento del MVP.",
    )
    add_key_value_table(
        doc,
        [
            ("lib/app", "Configuración general de la aplicación, rutas y tema."),
            ("lib/core", "Configuraciones, constantes, errores, red y utilidades comunes."),
            ("lib/shared", "Modelos y widgets reutilizables."),
            ("lib/features/auth", "Inicio de sesión y control de roles."),
            ("lib/features/admin", "Panel administrativo, registro y listado de usuarios."),
            ("lib/features/client", "Pantalla principal del usuario del gimnasio."),
            ("lib/features/memberships", "Mensualidades y estado de vigencia."),
            ("lib/features/exercises", "Ejercicios clasificados por grupo muscular."),
            ("lib/features/routines", "Rutina semanal del usuario."),
            ("lib/features/profile", "Perfil y datos del usuario."),
        ],
    )
    add_capture_box(doc, "Captura 9", "árbol de carpetas dentro de lib", 1.9)
    add_capture_box(doc, "Captura 10", "carpetas assets/images, assets/icons y assets/animations", 1.2)


def add_theme(doc):
    add_heading(doc, 7, "Configuración del tema inicial")
    add_paragraph(
        doc,
        "El tema inicial definirá la identidad visual de GymPro. Se recomienda utilizar una paleta relacionada con salud, energía y confianza, manteniendo una interfaz clara para usuarios no técnicos.",
    )
    add_bullet(doc, "Color primario sugerido: verde deportivo o azul de confianza.")
    add_bullet(doc, "Uso de Material Design para componentes móviles.")
    add_bullet(doc, "Botones, campos de texto y tarjetas con estilo consistente.")
    add_bullet(doc, "Diseño adaptable a pantallas de teléfonos móviles.")
    add_capture_box(doc, "Captura 11", "archivo o pantalla donde se configure el tema inicial", 1.8)
    add_capture_box(doc, "Captura 12", "vista preliminar de login con el tema aplicado", 1.8)


def add_navigation(doc):
    add_heading(doc, 8, "Flujo de navegación")
    add_paragraph(
        doc,
        "El flujo de navegación se basa en el acceso por roles. El administrador ingresa al panel administrativo y el usuario del gimnasio ingresa a las opciones de entrenamiento, rutina y perfil.",
    )
    add_step(doc, 1, "El usuario abre la aplicación móvil.")
    add_step(doc, 2, "La aplicación muestra la pantalla de inicio de sesión.")
    add_step(doc, 3, "El sistema valida correo y contraseña.")
    add_step(doc, 4, "Si el rol es administrador, se redirige al panel administrativo.")
    add_step(doc, 5, "Si el rol es cliente, se redirige al panel principal del usuario.")
    add_step(doc, 6, "El usuario puede cerrar sesión y volver al login.")
    add_key_value_table(
        doc,
        [
            ("Administrador", "Login -> Panel administrador -> Registro/Listado de usuarios -> Cerrar sesión"),
            ("Cliente", "Login -> Inicio cliente -> Entrenamiento/Rutina/Perfil -> Cerrar sesión"),
        ],
    )
    add_capture_box(doc, "Captura 13", "diagrama o wireframe del flujo de navegación", 1.8)
    add_capture_box(doc, "Captura 14", "pantallas principales según rol", 1.6)


def add_readme(doc):
    add_heading(doc, 9, "README del Proyecto")
    add_paragraph(
        doc,
        "El archivo README debe describir el propósito del proyecto, tecnologías utilizadas, pasos de instalación y forma de ejecución. Este contenido servirá como guía rápida para cualquier persona que revise el repositorio.",
    )
    add_heading(doc, "9.1", "Contenido sugerido del README", 2)
    add_bullet(doc, "Nombre del proyecto: GymPro App.")
    add_bullet(doc, "Descripción breve del problema que resuelve.")
    add_bullet(doc, "Tecnologías utilizadas: Flutter, Dart, Supabase/PostgreSQL y GitHub.")
    add_bullet(doc, "Instrucciones de instalación y ejecución.")
    add_bullet(doc, "Estructura de carpetas del proyecto.")
    add_bullet(doc, "Autores del proyecto.")
    add_capture_box(doc, "Captura 15", "README.md abierto en el editor", 1.8)
    add_capture_box(doc, "Captura 16", "README visible en GitHub", 1.6)


def add_conclusion(doc):
    add_heading(doc, 10, "Conclusiones")
    add_paragraph(
        doc,
        "La preparación inicial del proyecto GymPro establece una base ordenada para el desarrollo del MVP. La estructura creada permite separar autenticación, administración, membresías, ejercicios, rutinas y perfil, lo que facilitará la conexión futura con Supabase y la implementación de las pantallas móviles.",
    )
    add_bullet(doc, "El proyecto Flutter se encuentra inicializado y listo para organizar la lógica principal dentro de lib.")
    add_bullet(doc, "La estructura modular responde a los requerimientos funcionales definidos para el MVP.")
    add_bullet(doc, "Los espacios de evidencia permitirán documentar capturas reales del proceso de configuración y ejecución.")


def main():
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    doc.add_page_break()
    add_index(doc)
    doc.add_page_break()

    for section_builder in [
        add_environment,
        add_project_creation,
        add_execution,
        add_git,
        add_structure,
        add_theme,
        add_navigation,
        add_readme,
        add_conclusion,
    ]:
        section_builder(doc)
        if section_builder != add_conclusion:
            doc.add_page_break()

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
