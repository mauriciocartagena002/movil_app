from pathlib import Path

from docx import Document


DOCX = Path("docs/actividad_1_avance2_corregido.docx")


INDEX_LINES = [
    "Portada ............................................................ 1",
    "Configuración del entorno .......................................... 2",
    "Creación del proyecto Flutter ..................................... 4",
    "Ejecución de la aplicación ........................................ 5",
    "Configuración de Git y GitHub ..................................... 6",
    "Estructura inicial del proyecto ................................... 7",
    "Configuración del tema inicial .................................... 8",
    "Flujo de navegación ............................................... 10",
    "README del Proyecto ............................................... 11",
    "Conclusiones ...................................................... 11",
]


def main():
    doc = Document(DOCX)

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.text = paragraph.text.strip()
        if paragraph.text.strip() == "20":
            paragraph.text = "2026"

    index_start = None
    first_content = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "ÍNDICE DE CONTENIDO":
            index_start = i
        if index_start is not None and paragraph.text.strip() == "Configuración del entorno":
            first_content = i
            break

    if index_start is not None and first_content is not None:
        index_paragraphs = doc.paragraphs[index_start + 1:first_content]
        for paragraph, line in zip(index_paragraphs, INDEX_LINES):
            paragraph.text = line
            paragraph.style = doc.styles["Normal"]
        for paragraph in index_paragraphs[len(INDEX_LINES):]:
            paragraph.text = ""

    doc.save(DOCX)
    print(DOCX.resolve())


if __name__ == "__main__":
    main()
