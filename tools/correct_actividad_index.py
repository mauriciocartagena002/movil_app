from pathlib import Path

from docx import Document


SOURCE = Path(r"C:\Users\user\Downloads\actividad_1 avance2.docx")
OUT = Path("docs/actividad_1_avance2_corregido.docx")


INDEX_MARKER = "ÍNDICE DE CONTENIDO"


def replace_index_block(doc):
    index_start = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == INDEX_MARKER:
            index_start = i
            break

    if index_start is None:
        return

    new_lines = [
        "Portada ............................................................ 1",
        "Configuración del entorno .......................................... __",
        "Creación del proyecto Flutter ..................................... __",
        "Ejecución de la aplicación ........................................ __",
        "Configuración de Git y GitHub ..................................... __",
        "Estructura inicial del proyecto ................................... __",
        "Configuración del tema inicial .................................... __",
        "Flujo de navegación ............................................... __",
        "README del Proyecto ............................................... __",
        "Conclusiones ...................................................... __",
    ]

    first_content = None
    for i in range(index_start + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip() == "Configuración del entorno":
            first_content = i
            break

    if first_content is None:
        return

    existing = doc.paragraphs[index_start + 1:first_content]
    for paragraph, text in zip(existing, new_lines):
        paragraph.text = text
        paragraph.style = doc.styles["Normal"]

    for paragraph in existing[len(new_lines):]:
        paragraph.text = ""

    anchor = existing[-1] if existing else doc.paragraphs[index_start]
    if len(existing) < len(new_lines):
        for text in new_lines[len(existing):]:
            anchor = anchor.insert_paragraph_before(text, style="Normal")


def ensure_readme_section(doc):
    has_readme_body = any(
        p.text.strip() == "README del Proyecto" and p.style.name.startswith("Heading")
        for p in doc.paragraphs
    )
    if has_readme_body:
        return

    conclusion = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Conclusiones":
            conclusion = paragraph
            break

    if conclusion is None:
        return

    conclusion.insert_paragraph_before("", style="Normal")
    readme_heading = conclusion.insert_paragraph_before("README del Proyecto", style="Heading 1")
    readme_heading.paragraph_format.space_before = doc.styles["Heading 1"].paragraph_format.space_before
    conclusion.insert_paragraph_before(
        "El archivo README.md cumple la función de presentar la información básica del proyecto GymPro. "
        "En este documento se debe describir el objetivo de la aplicación, las tecnologías utilizadas, "
        "la forma de ejecutar el proyecto y una breve explicación de la estructura de carpetas.",
        style="Normal",
    )
    conclusion.insert_paragraph_before(
        "Además, el README servirá como guía inicial para otros desarrolladores o revisores del proyecto, "
        "permitiendo comprender rápidamente que la aplicación está desarrollada con Flutter y Dart, y que "
        "su propósito es apoyar la gestión de usuarios, membresías, rutinas y ejercicios de un gimnasio.",
        style="Normal",
    )


def fix_conclusions(doc):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Conclusiones":
            paragraph.style = doc.styles["Heading 1"]
        if "Queda pendiente documentar Git/GitHub" in paragraph.text:
            paragraph.text = (
                "Se documentó la configuración básica de Git y GitHub mediante la evidencia del repositorio remoto."
            )


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)
    replace_index_block(doc)
    ensure_readme_section(doc)
    fix_conclusions(doc)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
